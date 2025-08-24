import io
import logging
from typing import Dict, Any, List, Optional, Tuple, Union
from pathlib import Path
import hashlib
import mimetypes
from dataclasses import dataclass

# Document processing libraries
import PyPDF2
from PIL import Image
import docx
from docx.document import Document as DocxDocument

from app.core.exceptions import (
    DocumentProcessingError,
    DocumentTooLargeError,
    UnsupportedDocumentFormatError,
    DocumentCorruptedError
)
from app.config import settings

logger = logging.getLogger(__name__)


@dataclass
class DocumentContent:
    """Represents processed document content."""
    text: str
    images: List[Dict[str, Any]]
    tables: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    content_hash: str
    file_type: str
    file_size: int
    
    def get_total_content_length(self) -> int:
        """Get total length of all text content."""
        total = len(self.text)
        for table in self.tables:
            total += len(str(table.get('content', '')))
        return total


class DocumentProcessor:
    """Processes various document formats for AI analysis."""
    
    SUPPORTED_FORMATS = {
        'application/pdf': ['.pdf'],
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['.docx'],
        'application/msword': ['.doc'],
        'text/plain': ['.txt'],
        'image/png': ['.png'],
        'image/jpeg': ['.jpg', '.jpeg'],
        'image/gif': ['.gif'],
        'image/bmp': ['.bmp'],
        'image/tiff': ['.tiff', '.tif']
    }
    
    def __init__(self):
        self.max_file_size = settings.processing.max_file_size
        self.max_text_length = 2000000  # 2M characters for Gemini context
    
    def process_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: Optional[str] = None
    ) -> DocumentContent:
        """Process a document and extract all relevant content."""
        
        # Validate file size
        if len(file_content) > self.max_file_size:
            raise DocumentTooLargeError(len(file_content), self.max_file_size)
        
        # Determine file type
        file_type = self._detect_file_type(file_content, filename, content_type)
        
        # Generate content hash
        content_hash = hashlib.sha256(file_content).hexdigest()
        
        # Process based on file type
        try:
            if file_type == 'application/pdf':
                return self._process_pdf(file_content, filename, content_hash, file_type)
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return self._process_docx(file_content, filename, content_hash, file_type)
            elif file_type.startswith('image/'):
                return self._process_image(file_content, filename, content_hash, file_type)
            elif file_type == 'text/plain':
                return self._process_text(file_content, filename, content_hash, file_type)
            else:
                raise UnsupportedDocumentFormatError(
                    file_type, 
                    list(self.SUPPORTED_FORMATS.keys())
                )
                
        except Exception as e:
            if isinstance(e, (DocumentProcessingError, UnsupportedDocumentFormatError)):
                raise
            else:
                logger.error(f"Unexpected error processing document: {e}")
                raise DocumentCorruptedError(f"Failed to process document: {str(e)}")
    
    def _detect_file_type(
        self,
        content: bytes,
        filename: str,
        content_type: Optional[str] = None
    ) -> str:
        """Detect the actual file type."""
        
        # First try the provided content type
        if content_type and content_type in self.SUPPORTED_FORMATS:
            return content_type
        
        # Try to guess from filename extension
        guessed_type, _ = mimetypes.guess_type(filename)
        if guessed_type and guessed_type in self.SUPPORTED_FORMATS:
            return guessed_type
        
        # Try to detect from file content (magic bytes)
        if content.startswith(b'%PDF'):
            return 'application/pdf'
        elif content.startswith(b'PK'):  # ZIP-based formats like DOCX
            # Check for DOCX specifically
            if b'word/' in content[:1000] or b'docProps' in content[:1000]:
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif content.startswith((b'\xff\xd8\xff', b'\x89PNG')):
            if content.startswith(b'\xff\xd8\xff'):
                return 'image/jpeg'
            else:
                return 'image/png'
        
        # Fallback to text if it appears to be UTF-8 text
        try:
            content.decode('utf-8')
            return 'text/plain'
        except UnicodeDecodeError:
            pass
        
        raise UnsupportedDocumentFormatError(
            f"unknown (detected from {filename})",
            list(self.SUPPORTED_FORMATS.keys())
        )
    
    def _process_pdf(
        self,
        content: bytes,
        filename: str,
        content_hash: str,
        file_type: str
    ) -> DocumentContent:
        """Process PDF document."""
        try:
            pdf_stream = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_stream)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                # Try to decrypt with empty password
                if not pdf_reader.decrypt(""):
                    raise DocumentCorruptedError("PDF is password protected")
            
            # Extract text from all pages
            text_parts = []
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            full_text = "\n\n".join(text_parts)
            
            # Truncate if too long
            if len(full_text) > self.max_text_length:
                logger.warning(f"Document text truncated from {len(full_text)} to {self.max_text_length} characters")
                full_text = full_text[:self.max_text_length] + "\n\n[CONTENT TRUNCATED]"
            
            # Extract metadata
            metadata = {
                "page_count": page_count,
                "title": getattr(pdf_reader.metadata, 'title', None) if pdf_reader.metadata else None,
                "author": getattr(pdf_reader.metadata, 'author', None) if pdf_reader.metadata else None,
                "subject": getattr(pdf_reader.metadata, 'subject', None) if pdf_reader.metadata else None,
                "creator": getattr(pdf_reader.metadata, 'creator', None) if pdf_reader.metadata else None,
                "producer": getattr(pdf_reader.metadata, 'producer', None) if pdf_reader.metadata else None,
                "creation_date": getattr(pdf_reader.metadata, 'creation_date', None) if pdf_reader.metadata else None,
                "modification_date": getattr(pdf_reader.metadata, 'modification_date', None) if pdf_reader.metadata else None,
            }
            
            return DocumentContent(
                text=full_text,
                images=[],  # PDF image extraction would require additional libraries
                tables=[],  # Table extraction would require additional parsing
                metadata=metadata,
                content_hash=content_hash,
                file_type=file_type,
                file_size=len(content)
            )
            
        except Exception as e:
            if isinstance(e, DocumentCorruptedError):
                raise
            logger.error(f"Failed to process PDF: {e}")
            raise DocumentCorruptedError(f"PDF processing failed: {str(e)}")
    
    def _process_docx(
        self,
        content: bytes,
        filename: str,
        content_hash: str,
        file_type: str
    ) -> DocumentContent:
        """Process DOCX document."""
        try:
            doc_stream = io.BytesIO(content)
            doc = docx.Document(doc_stream)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            full_text = "\n".join(text_parts)
            
            # Extract tables
            tables = []
            for table_num, table in enumerate(doc.tables):
                table_data = {
                    "table_number": table_num + 1,
                    "rows": [],
                    "content": ""
                }
                
                table_text_parts = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                        table_text_parts.append(cell_text)
                    table_data["rows"].append(row_data)
                
                table_data["content"] = " | ".join(table_text_parts)
                tables.append(table_data)
            
            # Combine text and table content
            if tables:
                table_texts = [f"Table {t['table_number']}: {t['content']}" for t in tables]
                full_text += "\n\nTABLES:\n" + "\n".join(table_texts)
            
            # Truncate if too long
            if len(full_text) > self.max_text_length:
                logger.warning(f"Document text truncated from {len(full_text)} to {self.max_text_length} characters")
                full_text = full_text[:self.max_text_length] + "\n\n[CONTENT TRUNCATED]"
            
            # Extract metadata
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "title": getattr(doc.core_properties, 'title', None),
                "author": getattr(doc.core_properties, 'author', None),
                "subject": getattr(doc.core_properties, 'subject', None),
                "created": getattr(doc.core_properties, 'created', None),
                "modified": getattr(doc.core_properties, 'modified', None),
            }
            
            return DocumentContent(
                text=full_text,
                images=[],  # Image extraction would require additional processing
                tables=tables,
                metadata=metadata,
                content_hash=content_hash,
                file_type=file_type,
                file_size=len(content)
            )
            
        except Exception as e:
            logger.error(f"Failed to process DOCX: {e}")
            raise DocumentCorruptedError(f"DOCX processing failed: {str(e)}")
    
    def _process_image(
        self,
        content: bytes,
        filename: str,
        content_hash: str,
        file_type: str
    ) -> DocumentContent:
        """Process image document."""
        try:
            image_stream = io.BytesIO(content)
            image = Image.open(image_stream)
            
            # Get image metadata
            metadata = {
                "format": image.format,
                "mode": image.mode,
                "size": image.size,
                "width": image.width,
                "height": image.height,
                "info": dict(image.info) if hasattr(image, 'info') else {}
            }
            
            # For images, we return minimal text and rely on multimodal processing
            text = f"[IMAGE: {filename} - {image.width}x{image.height} pixels, {image.format} format]"
            
            # Store image data for multimodal processing
            images = [{
                "filename": filename,
                "format": image.format,
                "size": image.size,
                "mode": image.mode,
                "content_type": file_type,
                "data": content  # Include raw data for API calls
            }]
            
            return DocumentContent(
                text=text,
                images=images,
                tables=[],
                metadata=metadata,
                content_hash=content_hash,
                file_type=file_type,
                file_size=len(content)
            )
            
        except Exception as e:
            logger.error(f"Failed to process image: {e}")
            raise DocumentCorruptedError(f"Image processing failed: {str(e)}")
    
    def _process_text(
        self,
        content: bytes,
        filename: str,
        content_hash: str,
        file_type: str
    ) -> DocumentContent:
        """Process plain text document."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'iso-8859-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise DocumentCorruptedError("Could not decode text file with any common encoding")
            
            # Truncate if too long
            if len(text) > self.max_text_length:
                logger.warning(f"Text document truncated from {len(text)} to {self.max_text_length} characters")
                text = text[:self.max_text_length] + "\n\n[CONTENT TRUNCATED]"
            
            metadata = {
                "line_count": len(text.splitlines()),
                "character_count": len(text),
                "word_count": len(text.split())
            }
            
            return DocumentContent(
                text=text,
                images=[],
                tables=[],
                metadata=metadata,
                content_hash=content_hash,
                file_type=file_type,
                file_size=len(content)
            )
            
        except Exception as e:
            if isinstance(e, DocumentCorruptedError):
                raise
            logger.error(f"Failed to process text: {e}")
            raise DocumentCorruptedError(f"Text processing failed: {str(e)}")
    
    def get_multimodal_content(self, document: DocumentContent) -> List[Union[str, Dict[str, Any]]]:
        """Prepare content for multimodal AI processing."""
        content_parts = []
        
        # Add text content
        if document.text:
            content_parts.append(document.text)
        
        # Add images for multimodal processing
        for image in document.images:
            content_parts.append({
                "mime_type": image["content_type"],
                "data": image["data"]
            })
        
        return content_parts
    
    def validate_document(self, document: DocumentContent) -> Dict[str, Any]:
        """Validate processed document content."""
        validation = {
            "is_valid": True,
            "errors": [],
            "warnings": [],
            "metrics": {}
        }
        
        # Check if we have any extractable content
        total_content_length = document.get_total_content_length()
        if total_content_length == 0:
            validation["errors"].append("No extractable content found in document")
            validation["is_valid"] = False
        elif total_content_length < 100:
            validation["warnings"].append("Very little content extracted from document")
        
        # Check file size
        if document.file_size > self.max_file_size * 0.8:  # 80% of max
            validation["warnings"].append("Document is close to maximum file size limit")
        
        # Add metrics
        validation["metrics"] = {
            "file_size": document.file_size,
            "text_length": len(document.text),
            "image_count": len(document.images),
            "table_count": len(document.tables),
            "total_content_length": total_content_length,
            "content_hash": document.content_hash
        }
        
        return validation


# Global processor instance
_document_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    """Get the global document processor instance."""
    global _document_processor
    
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    
    return _document_processor


def reset_document_processor() -> None:
    """Reset the global document processor (useful for testing)."""
    global _document_processor
    _document_processor = None